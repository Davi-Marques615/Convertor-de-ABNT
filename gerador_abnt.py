from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE


@dataclass(frozen=True)
class ABNTConfig:
    font_family: str = "Arial"
    font_size_body: int = 12
    font_size_small: int = 10
    margin_top_front_mm: int = 30
    margin_left_front_mm: int = 30
    margin_bottom_front_mm: int = 20
    margin_right_front_mm: int = 20
    paragraph_first_line_indent_mm: float = 12.5
    long_quote_left_indent_mm: int = 40


class ABNTDocumentGenerator:
    def __init__(self, dados: dict[str, Any], caminho_saida: str | Path) -> None:
        self.dados = dados
        self.caminho_saida = Path(caminho_saida)
        self.config = ABNTConfig()
        self.documento = Document()
        self.paginas_pre_textuais_contadas = 0
        self.contadores_secao = [0] * 6

    def gerar(self) -> Path:
        self.caminho_saida.parent.mkdir(parents=True, exist_ok=True)
        self._configurar_documento()
        self._criar_elementos_pre_textuais()
        self._criar_elementos_textuais()
        self._criar_elementos_pos_textuais()
        self.documento.save(self.caminho_saida)
        return self.caminho_saida

    def _configurar_documento(self) -> None:
        self._configurar_secao(self.documento.sections[0])
        self._configurar_estilos()
        self._ativar_atualizacao_automatica_do_sumario()

    def _ativar_atualizacao_automatica_do_sumario(self) -> None:
        settings = self.documento.settings.element
        update_fields = settings.find(qn("w:updateFields"))
        if update_fields is None:
            update_fields = OxmlElement("w:updateFields")
            settings.append(update_fields)
        update_fields.set(qn("w:val"), "true")

    def _configurar_secao(self, secao: Any) -> None:
        secao.page_width = Mm(210)
        secao.page_height = Mm(297)
        secao.top_margin = Mm(self.config.margin_top_front_mm)
        secao.left_margin = Mm(self.config.margin_left_front_mm)
        secao.bottom_margin = Mm(self.config.margin_bottom_front_mm)
        secao.right_margin = Mm(self.config.margin_right_front_mm)
        secao.header_distance = Mm(20)

    def _configurar_estilos(self) -> None:
        normal = self.documento.styles["Normal"]
        normal.font.name = self.config.font_family
        normal.font.size = Pt(self.config.font_size_body)
        normal.font.color.rgb = RGBColor(0, 0, 0)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), self.config.font_family)
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal.paragraph_format.line_spacing = 1.5
        normal.paragraph_format.first_line_indent = Mm(self.config.paragraph_first_line_indent_mm)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(0)

        for nivel in range(1, 6):
            estilo = self.documento.styles[f"Heading {nivel}"]
            estilo.font.name = self.config.font_family
            estilo.font.size = Pt(self.config.font_size_body)
            estilo.font.bold = True
            estilo.font.color.rgb = RGBColor(0, 0, 0)
            estilo._element.rPr.rFonts.set(qn("w:eastAsia"), self.config.font_family)
            estilo.paragraph_format.line_spacing = 1.5
            estilo.paragraph_format.first_line_indent = None
            estilo.paragraph_format.space_before = Pt(12)
            estilo.paragraph_format.space_after = Pt(12)

    def _criar_elementos_pre_textuais(self) -> None:
        self._adicionar_capa()
        self._quebrar_pagina()

        self._adicionar_folha_rosto()
        self._contar_pagina_pre_textual()

        if self._deve_incluir_folha_aprovacao():
            self._quebrar_pagina()
            self._adicionar_folha_aprovacao()
            self._contar_pagina_pre_textual()

        self._adicionar_elemento_opcional_pre_textual("dedicatoria", "DEDICATÓRIA", alinhamento=WD_ALIGN_PARAGRAPH.RIGHT)
        self._adicionar_elemento_opcional_pre_textual("agradecimentos", "AGRADECIMENTOS")
        self._adicionar_elemento_opcional_pre_textual("epigrafe", "EPÍGRAFE", alinhamento=WD_ALIGN_PARAGRAPH.RIGHT)

        if self._texto("resumo"):
            self._quebrar_pagina()
            self._adicionar_resumo("RESUMO", self._texto("resumo"), self._lista("palavras_chave"), "Palavras-chave:")
            self._contar_pagina_pre_textual()

        if self._texto("abstract"):
            self._quebrar_pagina()
            self._adicionar_resumo("ABSTRACT", self._texto("abstract"), self._lista("keywords"), "Keywords:")
            self._contar_pagina_pre_textual()

        self._quebrar_pagina()
        self._adicionar_sumario()
        self._contar_pagina_pre_textual()

    def _adicionar_capa(self) -> None:
        instituicao = self._texto("instituicao")
        autores = self._lista("autores")
        titulo = self._texto("titulo")
        subtitulo = self._texto("subtitulo")
        cidade = self._texto("cidade")
        ano = self._texto("ano")

        if instituicao:
            self._adicionar_paragrafo(instituicao.upper(), alinhamento=WD_ALIGN_PARAGRAPH.CENTER, negrito=True, recuo=False)

        self._adicionar_espaco(3)
        for autor in autores:
            self._adicionar_paragrafo(autor.upper(), alinhamento=WD_ALIGN_PARAGRAPH.CENTER, recuo=False)

        self._adicionar_espaco(8)
        self._adicionar_paragrafo(titulo.upper(), alinhamento=WD_ALIGN_PARAGRAPH.CENTER, negrito=True, recuo=False)
        if subtitulo:
            self._adicionar_paragrafo(subtitulo, alinhamento=WD_ALIGN_PARAGRAPH.CENTER, negrito=True, recuo=False)

        self._adicionar_bloco_local_data(cidade, ano, espaco_antes_mm=84)

    def _adicionar_folha_rosto(self) -> None:
        autores = self._lista("autores")
        titulo = self._texto("titulo")
        subtitulo = self._texto("subtitulo")
        natureza = self._texto("natureza_trabalho") or self._montar_natureza_padrao()
        cidade = self._texto("cidade")
        ano = self._texto("ano")

        for autor in autores:
            self._adicionar_paragrafo(autor.upper(), alinhamento=WD_ALIGN_PARAGRAPH.CENTER, recuo=False)

        self._adicionar_espaco(7)
        self._adicionar_paragrafo(titulo.upper(), alinhamento=WD_ALIGN_PARAGRAPH.CENTER, negrito=True, recuo=False)
        if subtitulo:
            self._adicionar_paragrafo(subtitulo, alinhamento=WD_ALIGN_PARAGRAPH.CENTER, negrito=True, recuo=False)

        self._adicionar_espaco(4)
        paragrafo = self._adicionar_paragrafo(natureza, alinhamento=WD_ALIGN_PARAGRAPH.JUSTIFY, recuo=False)
        paragrafo.paragraph_format.left_indent = Mm(80)
        paragrafo.paragraph_format.line_spacing = 1.0

        self._adicionar_bloco_local_data(cidade, ano, espaco_antes_mm=58)

    def _adicionar_folha_aprovacao(self) -> None:
        self._adicionar_titulo_centralizado("FOLHA DE APROVAÇÃO")
        self._adicionar_espaco(2)
        self._adicionar_paragrafo(self._texto("titulo").upper(), alinhamento=WD_ALIGN_PARAGRAPH.CENTER, negrito=True, recuo=False)
        self._adicionar_espaco(2)
        self._adicionar_paragrafo(self._texto("texto_aprovacao") or self._montar_texto_aprovacao_padrao())
        self._adicionar_espaco(3)
        self._adicionar_paragrafo("BANCA EXAMINADORA", alinhamento=WD_ALIGN_PARAGRAPH.CENTER, negrito=True, recuo=False)
        self._adicionar_espaco(2)

        membros = self._lista("banca_examinadora") or [self._texto("professor")]
        for membro in [m for m in membros if m]:
            self._adicionar_linha_assinatura(membro)

    def _adicionar_elemento_opcional_pre_textual(
        self,
        chave: str,
        titulo: str,
        alinhamento: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
    ) -> None:
        conteudo = self._texto(chave)
        if not conteudo:
            return
        self._quebrar_pagina()
        self._adicionar_titulo_centralizado(titulo)
        self._adicionar_espaco(3)
        for paragrafo in self._separar_paragrafos(conteudo):
            self._adicionar_paragrafo(paragrafo, alinhamento=alinhamento)
        self._contar_pagina_pre_textual()

    def _adicionar_resumo(self, titulo: str, texto: str, palavras: list[str], prefixo: str) -> None:
        self._adicionar_titulo_centralizado(titulo)
        self._adicionar_espaco(1)
        paragrafo = self._adicionar_paragrafo(texto, alinhamento=WD_ALIGN_PARAGRAPH.JUSTIFY, recuo=False)
        paragrafo.paragraph_format.line_spacing = 1.0
        if palavras:
            texto_palavras = f"{prefixo} {'; '.join(palavras)}."
            self._adicionar_espaco(1)
            paragrafo_palavras = self._adicionar_paragrafo(texto_palavras, alinhamento=WD_ALIGN_PARAGRAPH.JUSTIFY, recuo=False)
            paragrafo_palavras.paragraph_format.line_spacing = 1.0

    def _adicionar_sumario(self) -> None:
        self._adicionar_titulo_centralizado("SUMÁRIO")
        self._adicionar_espaco(1)
        paragrafo = self.documento.add_paragraph()
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragrafo.paragraph_format.line_spacing = 1.5
        self._adicionar_campo(paragrafo, r'TOC \o "1-5" \h \z \u')

    def _criar_elementos_textuais(self) -> None:
        secao_textual = self.documento.add_section(WD_SECTION.NEW_PAGE)
        self._configurar_secao(secao_textual)
        self._configurar_paginacao_visivel(secao_textual, self.paginas_pre_textuais_contadas + 1)

        secoes = self._obter_secoes_textuais()
        for secao in secoes:
            nivel = secao.get("nivel", 1)
            numero_formatado = self._gerar_numero_secao(nivel)
            self._adicionar_titulo_numerado(numero_formatado, secao["titulo"], nivel=nivel)
            self._adicionar_conteudo_textual(secao.get("conteudo", ""))
            self._adicionar_imagens_da_secao(secao.get("imagens", []), numero_formatado)

    def _gerar_numero_secao(self, nivel: int) -> str:
        # Incrementa o nível atual e reseta os subníveis
        self.contadores_secao[nivel] += 1
        for i in range(nivel + 1, 6):
            self.contadores_secao[i] = 0
        
        # Monta a string de numeração (ex: 1.1.2)
        partes = [str(self.contadores_secao[i]) for i in range(1, nivel + 1)]
        return ".".join(partes)

    def _obter_secoes_textuais(self) -> list[dict[str, Any]]:
        secoes_informadas = self.dados.get("secoes_textuais")
        if isinstance(secoes_informadas, list) and secoes_informadas:
            return secoes_informadas

        secoes: list[dict[str, Any]] = []
        if self._texto("introducao"):
            secoes.append({"titulo": "INTRODUÇÃO", "conteudo": self._texto("introducao"), "nivel": 1})
        if self._texto("desenvolvimento"):
            secoes.append({"titulo": "DESENVOLVIMENTO", "conteudo": self._texto("desenvolvimento"), "nivel": 1})
        if self._texto("conclusao"):
            secoes.append({"titulo": "CONCLUSÃO", "conteudo": self._texto("conclusao"), "nivel": 1})

        texto_completo = self._texto("texto_completo")
        if not secoes and texto_completo:
            secoes.append({"titulo": "INTRODUÇÃO", "conteudo": texto_completo, "nivel": 1})
        return secoes or [{"titulo": "INTRODUÇÃO", "conteudo": "", "nivel": 1}]

    def _adicionar_conteudo_textual(self, conteudo: str) -> None:
        for bloco in self._separar_paragrafos(conteudo):
            if bloco.startswith(">>>"):
                self._adicionar_citacao_longa(bloco.removeprefix(">>>").strip())
            else:
                self._adicionar_paragrafo(bloco, alinhamento=WD_ALIGN_PARAGRAPH.JUSTIFY)

    def _adicionar_imagens_da_secao(self, imagens: Any, numero_secao: str) -> None:
        if not isinstance(imagens, list):
            return
        for indice, imagem in enumerate(imagens[:1], start=1):
            if not isinstance(imagem, dict) or not imagem.get("caminho"):
                continue
            alinhamentos = {"left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT, "center": WD_ALIGN_PARAGRAPH.CENTER}
            alinhamento = alinhamentos.get(str(imagem.get("alinhamento", "center")).lower(), WD_ALIGN_PARAGRAPH.CENTER)
            titulo = str(imagem.get("titulo", "")).strip()
            fonte = str(imagem.get("fonte", "")).strip()
            if titulo:
                legenda = self._adicionar_paragrafo(titulo, alinhamento=alinhamento, recuo=False, tamanho=self.config.font_size_small)
                legenda.paragraph_format.line_spacing = 1.0
            try:
                largura = max(20.0, min(float(imagem.get("largura_mm", 100)), 170.0))
                paragrafo = self.documento.add_paragraph()
                paragrafo.alignment = alinhamento
                paragrafo.paragraph_format.line_spacing = 1.0
                paragrafo.paragraph_format.first_line_indent = None
                paragrafo.add_run().add_picture(str(imagem["caminho"]), width=Mm(largura))
            except (OSError, ValueError, TypeError):
                continue
            if fonte:
                fonte_paragrafo = self._adicionar_paragrafo(fonte, alinhamento=WD_ALIGN_PARAGRAPH.LEFT, recuo=False, tamanho=self.config.font_size_small)
                fonte_paragrafo.paragraph_format.line_spacing = 1.0
            if indice < len(imagens[:1]):
                self._adicionar_espaco(1)

    def _criar_elementos_pos_textuais(self) -> None:
        referencias = self._lista("referencias")
        if referencias:
            self._quebrar_pagina()
            self._adicionar_referencias(referencias)

        self._adicionar_elemento_pos_textual_opcional("glossario", "GLOSSÁRIO")
        self._adicionar_elemento_pos_textual_opcional("apendice", "APÊNDICE A - MATERIAL COMPLEMENTAR")
        self._adicionar_elemento_pos_textual_opcional("anexo", "ANEXO A - MATERIAL COMPLEMENTAR")
        self._adicionar_elemento_pos_textual_opcional("indice", "ÍNDICE")

    def _adicionar_referencias(self, referencias: list[str]) -> None:
        self._adicionar_titulo_centralizado("REFERÊNCIAS")
        self._adicionar_espaco(1)
        for referencia in sorted(referencias, key=lambda item: item.upper()):
            paragrafo = self._adicionar_paragrafo(referencia, alinhamento=WD_ALIGN_PARAGRAPH.LEFT, recuo=False)
            paragrafo.paragraph_format.line_spacing = 1.0
            self._adicionar_espaco(1)

    def _adicionar_elemento_pos_textual_opcional(self, chave: str, titulo: str) -> None:
        conteudo = self._texto(chave)
        if not conteudo:
            return
        self._quebrar_pagina()
        self._adicionar_titulo_centralizado(titulo)
        self._adicionar_espaco(1)
        for paragrafo in self._separar_paragrafos(conteudo):
            self._adicionar_paragrafo(paragrafo)

    def _adicionar_titulo_numerado(self, numero: str, titulo: str, nivel: int) -> None:
        texto = f"{numero} {titulo.upper()}" if nivel == 1 else f"{numero} {titulo}"
        paragrafo = self.documento.add_paragraph(style=f"Heading {nivel}")
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragrafo.paragraph_format.first_line_indent = None
        run = paragrafo.add_run(texto)
        self._formatar_run(run, negrito=True)

    def _adicionar_titulo_centralizado(self, texto: str) -> None:
        paragrafo = self._adicionar_paragrafo(texto.upper(), alinhamento=WD_ALIGN_PARAGRAPH.CENTER, negrito=True, recuo=False)
        paragrafo.paragraph_format.line_spacing = 1.5

    def _adicionar_citacao_longa(self, texto: str) -> None:
        paragrafo = self._adicionar_paragrafo(texto, alinhamento=WD_ALIGN_PARAGRAPH.JUSTIFY, tamanho=self.config.font_size_small, recuo=False)
        paragrafo.paragraph_format.left_indent = Mm(self.config.long_quote_left_indent_mm)
        paragrafo.paragraph_format.line_spacing = 1.0

    def _adicionar_linha_assinatura(self, nome: str) -> None:
        self._adicionar_paragrafo("________________________________________", alinhamento=WD_ALIGN_PARAGRAPH.CENTER, recuo=False)
        self._adicionar_paragrafo(nome, alinhamento=WD_ALIGN_PARAGRAPH.CENTER, recuo=False)
        self._adicionar_espaco(2)

    def _adicionar_bloco_local_data(self, cidade: str, ano: str, espaco_antes_mm: int) -> None:
        paragrafo_cidade = self._adicionar_paragrafo(cidade.upper(), alinhamento=WD_ALIGN_PARAGRAPH.CENTER, recuo=False)
        paragrafo_cidade.paragraph_format.space_before = Mm(espaco_antes_mm)
        paragrafo_cidade.paragraph_format.keep_with_next = True
        self._adicionar_paragrafo(ano, alinhamento=WD_ALIGN_PARAGRAPH.CENTER, recuo=False)

    def _adicionar_paragrafo(
        self,
        texto: str,
        alinhamento: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
        negrito: bool = False,
        tamanho: int | None = None,
        recuo: bool = True,
    ) -> Any:
        paragrafo = self.documento.add_paragraph()
        paragrafo.alignment = alinhamento
        paragrafo.paragraph_format.line_spacing = 1.5
        paragrafo.paragraph_format.space_before = Pt(0)
        paragrafo.paragraph_format.space_after = Pt(0)
        paragrafo.paragraph_format.first_line_indent = Mm(self.config.paragraph_first_line_indent_mm) if recuo else None
        run = paragrafo.add_run(texto)
        self._formatar_run(run, negrito=negrito, tamanho=tamanho)
        return paragrafo

    def _formatar_run(self, run: Any, negrito: bool = False, tamanho: int | None = None) -> None:
        run.font.name = self.config.font_family
        run.font.size = Pt(tamanho or self.config.font_size_body)
        run.font.bold = negrito
        run.font.color.rgb = RGBColor(0, 0, 0)

    def _adicionar_espaco(self, quantidade: int) -> None:
        for _ in range(quantidade):
            paragrafo = self.documento.add_paragraph()
            paragrafo.paragraph_format.line_spacing = 1.5
            paragrafo.paragraph_format.space_before = Pt(0)
            paragrafo.paragraph_format.space_after = Pt(0)

    def _quebrar_pagina(self) -> None:
        self.documento.add_page_break()

    def _contar_pagina_pre_textual(self) -> None:
        self.paginas_pre_textuais_contadas += 1

    def _configurar_paginacao_visivel(self, secao: Any, inicio: int) -> None:
        secao.footer_distance = Mm(20)
        secao.header_distance = Mm(20)
        secao.start_type = WD_SECTION.NEW_PAGE
        
        # Ativa a numeração de páginas no cabeçalho (ABNT)
        header = secao.header
        paragrafo = header.paragraphs[0]
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Adiciona o campo de número de página
        run = paragrafo.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run._element.append(fldChar)
        run._element.append(instrText)
        run._element.append(fldChar2)
        run._element.append(fldChar3)
        
        # Define o número inicial da seção
        pgNum = OxmlElement('w:pgNumType')
        pgNum.set(qn('w:start'), str(inicio))
        secao._sectPr.append(pgNum)

    def _adicionar_campo(self, paragrafo: Any, texto: str) -> None:
        run = paragrafo.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = texto
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "end")
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)
        run._element.append(fldChar3)

    def _texto(self, chave: str) -> str:
        return str(self.dados.get(chave, "")).strip()

    def _lista(self, chave: str) -> list[str]:
        valor = self.dados.get(chave)
        if isinstance(valor, list):
            return [str(v).strip() for v in valor if str(v).strip()]
        return []

    def _separar_paragrafos(self, texto: str) -> list[str]:
        return [p.strip() for p in texto.split("\n") if p.strip()]

    def _deve_incluir_folha_aprovacao(self) -> bool:
        valor = self._texto("folha_aprovacao").lower()
        if valor == "sim":
            return True
        if valor == "nao":
            return False
        # Automático: inclui se houver banca examinadora ou professor
        return bool(self._lista("banca_examinadora") or self._texto("professor"))

    def _montar_natureza_padrao(self) -> str:
        tipo = self._texto("tipo_trabalho") or "Trabalho acadêmico"
        instituicao = self._texto("instituicao")
        curso = self._texto("curso_disciplina")
        return f"{tipo} apresentado à {instituicao} como requisito parcial para obtenção de nota na disciplina/curso de {curso}."

    def _montar_texto_aprovacao_padrao(self) -> str:
        return "Trabalho aprovado. A estrutura e o conteúdo estão em conformidade com as normas ABNT vigentes."


def gerar_documento_abnt(dados: dict[str, Any], caminho_saida: str | Path) -> Path:
    gerador = ABNTDocumentGenerator(dados, caminho_saida)
    return gerador.gerar()
